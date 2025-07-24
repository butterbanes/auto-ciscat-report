from docx.document import Document as DocumentType
from docx import Document as DocumentInit
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from clean_data import clean_outcome_tables
from generate import gen_timestamp, auto_gen_name 
import pandas as pd
from typing import List, Dict
import os
from parse_html import get_newest_html


def set_cell_border(cell, **kwargs):
    """
    Set the cell border style of a argumented cell using argumented cell-border-locations in **kwargs (keyword arguments)

    Parameters:
    cell (docx.Document.table.rows.cells): the individual cell being passed in from a loop inside add_df_to_docx()

    Returns:
    void: simply modifies the external word document table cell
    """

    # grab the cell's element
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # iterate over possible border locations
    for border_name in ['top', 'bottom', 'left', 'right', 'end', 'insideH', 'insideV']:
        # relate those border_names with passed arguments from **kwargs
        border = kwargs.get(border_name)
        if border:
            # if border is valid add an xml element with the format of the passed border_name
            border_el = OxmlElement(f"w:{border_name}")
            # iterate over the following keys to set the style of the border
            for key in ['sz', 'val', 'color', 'space']:
                if key in border:
                    border_el.set(qn(f"w:{key}"), str(border[key]))
            # add the xml border element to the tcPr
            tcPr.append(border_el)


def add_table(doc: DocumentType, dataframe: pd.DataFrame | str):
    table = None
    data: List[Dict] = [{}]
    if isinstance(dataframe, str):
        data = [{'0': 'NO TABLE DATA'}]
    else:
        data: List[Dict] = dataframe.to_dict('records')
    table = doc.add_table(rows=1, cols=len(data[0]))
    header_cells = table.rows[0].cells
    keys = data[0].keys()

    # add header row
    for idx, key in enumerate(keys):
        header_cells[idx].text = key

    # add data rows
    for item in data:
        row_cells = table.add_row().cells
        for idx, key in enumerate(keys):
            row_cells[idx].text = str(item[key])

# set font and border style
    for row in table.rows:
        for cell in row.cells:
            for pgraph in cell.paragraphs:
                for run in pgraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(7)
                set_cell_border(
                                cell, 
                                top={"sz": 12, "val": "single", "color": "007ACC", "space": "0"},
                                bottom={"sz": 12, "val": "single", "color": "007ACC", "space": "0"},
                                left={"sz": 12, "val": "single", "color": "007ACC", "space": "0"},
                                right={"sz": 12, "val": "single", "color": "007ACC", "space": "0"}
                            )

    


    # add paragraph break after the table
    doc.add_paragraph()


def add_horizontal_line(doc):
    """
    Add/print a horizontal line in a Word .docx file

    Parameters:
    doc (docx.Document): the Word document to manipulate

    Return:
    void: prints "behind-the-scenes' to the Word document and finishes
    """
    pgraph = doc.add_paragraph()
    pgraphPr = pgraph._element.get_or_add_pPr()
    pgraphBdr = OxmlElement('w:Bdr')
    
    # create an xml element and modify its attributes
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    # add the bottom to the paragraph border 
    # and then the paragraph border to the actual paragraph
    pgraphBdr.append(bottom)
    pgraphPr.append(pgraphBdr)


def add_outcome(doc: DocumentType, outcome: str):
    doc.add_paragraph(outcome)


def docx_writer(outcome_list, outcome_tables, new_document_name: str = get_newest_html()):
    doc = DocumentInit()
    doc.add_heading("REPORT", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f"{os.path.splitext(new_document_name)[0][:53]}", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f"{gen_timestamp()}", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for outcome, table in zip(outcome_list, outcome_tables):
        add_outcome(doc, outcome)
        add_table(doc, table)
        add_horizontal_line(doc)

    new_file = f'{auto_gen_name(new_document_name)}.docx'
    doc.save(f"{new_file}")
    return new_file 
